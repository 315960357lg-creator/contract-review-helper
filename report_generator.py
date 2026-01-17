"""
报告生成模块
支持导出为Word、PDF和Markdown格式
"""
import os
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    print("警告: python-docx 未安装")

from config import Config

logger = logging.getLogger(__name__)


class ReportGenerator:
    """报告生成器基类"""

    def __init__(self):
        self.output_dir = Config.OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        review_result: str,
        metadata: Dict[str, Any],
        output_format: str = "word"
    ) -> str:
        """
        生成报告

        Args:
            review_result: AI审查结果文本
            metadata: 元数据（合同信息、审查参数等）
            output_format: 输出格式 (word, pdf, markdown)

        Returns:
            生成的文件路径
        """
        raise NotImplementedError


class WordReportGenerator(ReportGenerator):
    """Word报告生成器"""

    def __init__(self):
        super().__init__()
        if Document is None:
            raise ImportError("python-docx 未安装")

    def generate(
        self,
        review_result: str,
        metadata: Dict[str, Any],
        output_format: str = "word"
    ) -> str:
        """生成Word格式报告"""
        try:
            doc = Document()

            # 设置文档字体（中文字体）
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 1. 标题
            title = doc.add_heading('合同审查报告', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 2. 基本信息
            doc.add_heading('一、基本信息', level=1)

            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Table Grid'

            info_data = [
                ('合同名称', metadata.get('contract_name', '未知')),
                ('客户身份', metadata.get('client_role', '未知')),
                ('合同类型', metadata.get('contract_type', '未知')),
                ('审查时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('审查要点数量', str(len(metadata.get('checklist', {}).get('specific_checks', []))))
            ]

            for i, (key, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = key
                info_table.rows[i].cells[1].text = value

            # 3. 审查要点
            doc.add_heading('二、审查要点清单', level=1)

            checklist = metadata.get('checklist', {})
            contract_focus = checklist.get('contract_focus', [])
            specific_checks = checklist.get('specific_checks', [])

            # 合同维度
            doc.add_paragraph('【审查维度】', style='List Bullet')
            for focus in contract_focus:
                doc.add_paragraph(f'• {focus}', style='List Bullet 2')

            # 具体审查点
            doc.add_paragraph('【具体审查点】', style='List Bullet')
            for check in specific_checks:
                p = doc.add_paragraph(f'• {check["point"]}', style='List Bullet 2')
                p.add_run(f'\n  {check["logic"]}').font.size = Pt(9)

            # 4. 审查结果（Markdown格式转Word）
            doc.add_heading('三、审查结果', level=1)

            # 简单的Markdown解析
            self._parse_markdown_to_word(doc, review_result)

            # 5. 免责声明
            doc.add_page_break()
            doc.add_heading('免责声明', level=2)
            disclaimer = doc.add_paragraph(
                '本报告由AI助手生成，仅供参考，不构成法律意见。'
                '请在签署任何法律文件前咨询专业律师。'
                '本助手不对使用本报告造成的任何后果承担责任。'
            )
            disclaimer.runs[0].font.size = Pt(9)
            disclaimer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # 保存文件
            filename = f"合同审查报告_{metadata.get('contract_name', '未知')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.output_dir / filename
            doc.save(str(filepath))

            logger.info(f"Word报告生成成功: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"生成Word报告失败: {e}")
            raise

    def _parse_markdown_to_word(self, doc: Document, markdown_text: str):
        """简单的Markdown转Word解析器"""
        lines = markdown_text.split('\n')
        in_table = False

        for line in lines:
            line = line.rstrip()

            if not line:
                doc.add_paragraph()
                continue

            # 一级标题
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            # 二级标题
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # 三级标题
            elif line.startswith('##### '):
                doc.add_heading(line[6:], level=5)
            # 列表
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # 表格（简单处理）
            elif line.startswith('|'):
                if not in_table:
                    in_table = True
                    # 创建表格
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    # 表头
                    for i, cell in enumerate(cells):
                        table.rows[0].cells[i].text = cell
                elif not line.startswith('|---'):
                    # 表格内容行
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    row = table.add_row()
                    for i, cell in enumerate(cells):
                        row.cells[i].text = cell
            else:
                # 普通段落（处理加粗）
                if '**' in line:
                    p = doc.add_paragraph()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            p.add_run(part)
                        else:
                            run = p.add_run(part)
                            run.bold = True
                else:
                    doc.add_paragraph(line)


class MarkdownReportGenerator(ReportGenerator):
    """Markdown报告生成器"""

    def generate(
        self,
        review_result: str,
        metadata: Dict[str, Any],
        output_format: str = "markdown"
    ) -> str:
        """生成Markdown格式报告"""
        try:
            # 构建Markdown内容
            content = f"""# 合同审查报告

## 一、基本信息

- **合同名称**: {metadata.get('contract_name', '未知')}
- **客户身份**: {metadata.get('client_role', '未知')}
- **合同类型**: {metadata.get('contract_type', '未知')}
- **审查时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 二、审查要点清单

### 审查维度
"""

            checklist = metadata.get('checklist', {})
            contract_focus = checklist.get('contract_focus', [])
            specific_checks = checklist.get('specific_checks', [])

            for focus in contract_focus:
                content += f"- {focus}\n"

            content += "\n### 具体审查点\n"
            for check in specific_checks:
                content += f"- **{check['point']}**: {check['logic']}\n"

            content += "\n## 三、审查结果\n\n"
            content += review_result

            content += """

---

## 免责声明

本报告由AI助手生成，仅供参考，不构成法律意见。请在签署任何法律文件前咨询专业律师。本助手不对使用本报告造成的任何后果承担责任。
"""

            # 保存文件
            filename = f"合同审查报告_{metadata.get('contract_name', '未知')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            filepath = self.output_dir / filename

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)

            logger.info(f"Markdown报告生成成功: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"生成Markdown报告失败: {e}")
            raise


class ReportGeneratorFactory:
    """报告生成器工厂类"""

    @staticmethod
    def create_generator(output_format: str) -> ReportGenerator:
        """
        创建报告生成器

        Args:
            output_format: 输出格式 (word, pdf, markdown)

        Returns:
            对应的报告生成器实例
        """
        if output_format in ["word", "docx"]:
            return WordReportGenerator()
        elif output_format in ["markdown", "md"]:
            return MarkdownReportGenerator()
        else:
            raise ValueError(f"不支持的输出格式: {output_format}")

    @staticmethod
    def generate_report(
        review_result: str,
        metadata: Dict[str, Any],
        output_format: str = "word"
    ) -> str:
        """
        生成报告（统一接口）

        Args:
            review_result: AI审查结果
            metadata: 元数据
            output_format: 输出格式

        Returns:
            生成的文件路径
        """
        generator = ReportGeneratorFactory.create_generator(output_format)
        return generator.generate(review_result, metadata, output_format)


# 测试代码
if __name__ == "__main__":
    logging.basicConfig(level=Config.LOG_LEVEL)

    # 测试报告生成
    test_metadata = {
        'contract_name': '测试合同',
        'client_role': '乙方',
        'contract_type': '软件开发合同',
        'checklist': {
            'contract_focus': ['付款条款', '知识产权', '违约责任'],
            'specific_checks': [
                {'point': '付款周期', 'logic': '检查付款周期是否合理'},
                {'point': '知识产权归属', 'logic': '确认知识产权归属'}
            ]
        }
    }

    test_review_result = """
### 一、核心风险提示

#### 风险点1: 付款周期过长
- **条款引用**: 第三条第2款
- **风险等级**: 高
- **风险分析**: 约定验收后90天付款，远超行业标准

### 二、修改方案对比

| 原条款内容 | 风险说明 | 修改建议 | 修改后建议文本 |
| :--- | :--- | :--- | :--- |
| 甲方应在验收合格后90日内支付款项 | 付款周期过长 | 缩短为30天 | 甲方应在验收合格后30日内支付款项 |
"""

    try:
        # 生成Word报告
        word_path = ReportGeneratorFactory.generate_report(
            test_review_result,
            test_metadata,
            "word"
        )
        print(f"Word报告生成成功: {word_path}")

        # 生成Markdown报告
        md_path = ReportGeneratorFactory.generate_report(
            test_review_result,
            test_metadata,
            "markdown"
        )
        print(f"Markdown报告生成成功: {md_path}")

    except Exception as e:
        print(f"测试失败: {e}")
        logger.exception(e)
