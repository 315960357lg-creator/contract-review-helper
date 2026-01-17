"""
文档解析模块
支持Word和PDF文档的解析
"""
from pathlib import Path
from typing import Dict, List, Optional
import logging

try:
    from docx import Document
except ImportError:
    Document = None
    print("警告: python-docx 未安装，Word文档解析功能不可用")

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
    print("警告: PyMuPDF 未安装，PDF文档解析功能不可用")

from config import Config

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器基类"""

    def __init__(self):
        self.text = ""
        self.metadata = {}

    def parse(self, file_path: str) -> Dict[str, any]:
        """
        解析文档

        Args:
            file_path: 文档路径

        Returns:
            包含文本和元数据的字典
        """
        raise NotImplementedError


class WordParser(DocumentParser):
    """Word文档解析器"""

    def __init__(self):
        super().__init__()
        if Document is None:
            raise ImportError("python-docx 未安装")

    def parse(self, file_path: str) -> Dict[str, any]:
        """
        解析Word文档

        Args:
            file_path: Word文档路径 (.docx)

        Returns:
            包含文本和元数据的字典
        """
        try:
            doc = Document(file_path)

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格文本
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            self.text = "\n\n".join(paragraphs)
            self.metadata = {
                "type": "word",
                "paragraphs_count": len(paragraphs),
                "tables_count": len(tables),
                "tables": tables
            }

            logger.info(f"成功解析Word文档: {file_path}")
            logger.info(f"提取段落数: {len(paragraphs)}, 表格数: {len(tables)}")

            return {
                "text": self.text,
                "metadata": self.metadata,
                "paragraphs": paragraphs,
                "tables": tables
            }

        except Exception as e:
            logger.error(f"解析Word文档失败: {e}")
            raise


class PDFParser(DocumentParser):
    """PDF文档解析器"""

    def __init__(self):
        super().__init__()
        if fitz is None:
            raise ImportError("PyMuPDF 未安装")

    def parse(self, file_path: str) -> Dict[str, any]:
        """
        解析PDF文档

        Args:
            file_path: PDF文档路径 (.pdf)

        Returns:
            包含文本和元数据的字典
        """
        try:
            doc = fitz.open(file_path)

            # 提取所有页面的文本
            text_blocks = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("blocks")
                for block in blocks:
                    if block[6] == 0:  # 文本块
                        text_blocks.append(block[4])

            self.text = "\n\n".join(text_blocks)
            self.metadata = {
                "type": "pdf",
                "pages": len(doc),
                "blocks_count": len(text_blocks)
            }

            doc.close()

            logger.info(f"成功解析PDF文档: {file_path}")
            logger.info(f"提取页数: {len(doc)}, 文本块数: {len(text_blocks)}")

            return {
                "text": self.text,
                "metadata": self.metadata,
                "blocks": text_blocks
            }

        except Exception as e:
            logger.error(f"解析PDF文档失败: {e}")
            raise


class DocumentParserFactory:
    """文档解析器工厂类"""

    @staticmethod
    def get_parser(file_path: str) -> DocumentParser:
        """
        根据文件类型返回对应的解析器

        Args:
            file_path: 文件路径

        Returns:
            对应的文档解析器实例

        Raises:
            ValueError: 不支持的文件类型
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return WordParser()
        elif suffix == ".pdf":
            return PDFParser()
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")

    @staticmethod
    def parse_document(file_path: str) -> Dict[str, any]:
        """
        解析文档（自动识别类型）

        Args:
            file_path: 文件路径

        Returns:
            解析结果字典
        """
        parser = DocumentParserFactory.get_parser(file_path)
        return parser.parse(file_path)


def extract_contract_sections(text: str) -> Dict[str, List[str]]:
    """
    从合同文本中提取章节

    Args:
        text: 合同全文

    Returns:
        章节字典
    """
    sections = {
        "general": [],  # 通用条款
        "core": [],  # 核心条款
        "other": []  # 其他
    }

    lines = text.split("\n")

    current_section = "other"
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 简单的章节识别逻辑（可以根据需要扩展）
        if any(keyword in line for keyword in ["第一条", "第二章", "1.", "一、"]):
            if any(keyword in line for keyword in ["违约", "责任", "赔偿", "付款"]):
                current_section = "core"
            else:
                current_section = "general"

        sections[current_section].append(line)

    return sections


# 测试代码
if __name__ == "__main__":
    logging.basicConfig(level=Config.LOG_LEVEL)

    # 测试文档解析
    test_file = "test.docx"  # 替换为实际文件路径

    try:
        result = DocumentParserFactory.parse_document(test_file)
        print(f"解析成功!")
        print(f"文本长度: {len(result['text'])}")
        print(f"元数据: {result['metadata']}")

        # 测试章节提取
        sections = extract_contract_sections(result['text'])
        print(f"\n章节统计:")
        print(f"通用条款: {len(sections['general'])} 行")
        print(f"核心条款: {len(sections['core'])} 行")

    except Exception as e:
        print(f"测试失败: {e}")
